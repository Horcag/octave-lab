#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для конвертации CSV файла с результатами в DOCX таблицу
"""

import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import os

def format_number(num):
    """
    Форматирует число для таблицы:
    - Обычные числа с 8 знаками после точки
    - Маленькие числа в экспоненциальном формате
    """
    if isinstance(num, str):
        return num
    
    if abs(num) < 1e-3 and num != 0:
        return "{:.8e}".format(num)
    else:
        return "{:.8f}".format(num)

def csv_to_docx(csv_file, output_docx, title="Таблица 4.1. Результаты расчетов"):
    """
    Конвертирует CSV файл в таблицу DOCX
    """
    # Чтение CSV файла
    df = pd.read_csv(csv_file)
    
    # Создание нового документа
    doc = Document()
    
    # Добавление заголовка
    heading = doc.add_paragraph()
    heading_run = heading.add_run(title)
    heading_run.bold = True
    heading_run.font.size = Pt(14)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Пустая строка
    doc.add_paragraph()
    
    # Создание таблицы
    num_rows, num_cols = df.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols)
    table.style = 'Table Grid'
    
    # Заполнение заголовков таблицы
    headers = df.columns.tolist()
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = str(header)
        cell_para = cell.paragraphs[0]
        cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell_para.runs:
            run.bold = True
    
    # Заполнение данных таблицы
    for row in range(num_rows):
        row_cells = table.rows[row + 1].cells
        for col in range(num_cols):
            value = df.iloc[row, col]
            row_cells[col].text = format_number(value)
            row_cells[col].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Настройка ширины столбцов таблицы
    table.autofit = False
    for column in table.columns:
        column.width = Cm(2.5)  # 2.5 см для всех столбцов
    
    # Сохранение документа
    doc.save(output_docx)
    print(f"Файл {output_docx} успешно создан!")

if __name__ == "__main__":
    # Путь к текущему каталогу, где находится скрипт
    current_dir = os.path.dirname(os.path.abspath(__file__))
    
    # Пути к файлам
    csv_file = os.path.join(current_dir, "results_table.csv")
    output_docx = os.path.join(current_dir, "results_table.docx")
    
    # Конвертация из CSV в DOCX
    csv_to_docx(csv_file, output_docx)
